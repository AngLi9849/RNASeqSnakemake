import sys

# logging
sys.stderr = open(snakemake.log[0], "w")

import os
from snakemake.shell import shell

shell(
    """
    Rscript workflow/scripts/R/docx.R
    """
)

# Create word template using the configured font
# Using script in Utkarsh Dalal's Answer in https://stackoverflow.com/questions/56658872/add-page-number-using-python-docx to set centred page numbers
# Using script in first section of https://python-docx.readthedocs.io/en/latest/user/styles-using.html to loop through styles and change font

import docx as docx
from docx.oxml import OxmlElement, ns
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

import matplotlib as matp

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)
    page_num_run = paragraph.add_run()
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)
    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')
    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"
    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')
    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)

doc = Document('resources/templates/r_template.docx')
#doc = Document()

styles = [
    s for s in doc.styles
#    s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH
]
for style in styles:
    if hasattr(doc.styles[str(style.name)],'font'):
        doc.styles[str(style.name)].font.name = snakemake.wildcards.font
        doc.styles[str(style.name)].font.color.rgb = docx.shared.RGBColor.from_string(matp.colors.cnames[str(snakemake.wildcards.font_colour)][1:])

#doc.styles.add_style('centered', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
#doc.styles['centered'].base_style = doc.styles['Normal']
#doc.styles['centered'].paragraph_format.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

doc.styles.add_style('heading 4', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
doc.styles['heading 4'].base_style = doc.styles['heading 3']

add_page_number(doc.sections[0].footer.paragraphs[0])
doc.save(snakemake.output.docx)

